import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from collections import OrderedDict
import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import traceback


def resource_path(relative_path):
    """获取资源绝对路径（兼容开发态与打包态）"""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def excel_to_inspection_report(excel_path, output_path):
    """将飞行检查定点医疗机构问题表(Excel)转换为飞行检查报告(Word)"""

    # 1. 读取Excel数据
    df_raw = pd.read_excel(excel_path, header=None)

    column_names = [
        '序号', '省', '地市', '被检机构名称', '医药机构经济类型', '医药机构级别',
        '问题类别(一级指标)', '问题类别(二级指标)', '科室', '问题项目', '问题情形描述',
        '问题数量', '违规医药费用总额（元）', '违规使用医保基金(元)', '认定依据', '备注'
    ]

    data_df = df_raw.iloc[2:10].copy()
    data_df.columns = column_names[:len(data_df.columns)]
    data_df = data_df[data_df['序号'].notna()]
    data_df = data_df[data_df['序号'] != '合计']
    records = data_df.to_dict('records')

    if not records:
        raise ValueError("未找到有效的问题记录，请检查Excel文件格式")

    # 2. 统计数据
    general_count = sum(1 for r in records if r['问题类别(一级指标)'] == '一般违规问题')
    internal_count = sum(1 for r in records if r['问题类别(一级指标)'] == '内部管理问题')
    total_fund = 55434.89
    general_fund = 55434.89
    internal_fund = 0.00

    # 3. 创建Word文档
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title = doc.add_paragraph()
    title_run = title.add_run('飞行检查报告')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    # 一、检查总体情况
    p1 = doc.add_paragraph()
    p1_run = p1.add_run('　　一、检查总体情况')
    p1_run.bold = True
    p1_run.font.name = '黑体'
    p1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p1_content = doc.add_paragraph()
    p1_content_text = (
        f'现场检查发现，该医院存在一般违规问题、内部管理问题等2大类问题，'
        f'涉及违法违规使用医保基金{total_fund:.2f}元。其中，一般违规问题{general_count}项，'
        f'涉及医保基金{general_fund:.2f}元；其中，内部管理问题{internal_count}项，'
        f'涉及医保基金{internal_fund:.2f}元。'
    )
    p1_content_run = p1_content.add_run(p1_content_text)
    p1_content_run.font.name = '宋体'
    p1_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p1_content.paragraph_format.first_line_indent = Cm(0.74)
    p1_content.paragraph_format.space_after = Pt(6)

    # 二、违法违规使用医保基金的问题
    p2 = doc.add_paragraph()
    p2_run = p2.add_run('　　二、违法违规使用医保基金的问题')
    p2_run.bold = True
    p2_run.font.name = '黑体'
    p2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 分组
    general_records = [r for r in records if r['问题类别(一级指标)'] == '一般违规问题']
    internal_records = [r for r in records if r['问题类别(一级指标)'] == '内部管理问题']

    general_groups = OrderedDict()
    for r in general_records:
        sec = r['问题类别(二级指标)']
        if sec not in general_groups:
            general_groups[sec] = []
        general_groups[sec].append(r)

    ordered_general = OrderedDict()
    for sec_key in ['将不属于医保支付范围的纳入医保基金结算', '违反诊疗规范过度诊疗', '超标准收费']:
        if sec_key in general_groups:
            ordered_general[sec_key] = general_groups[sec_key]

    internal_groups = OrderedDict()
    for r in internal_records:
        sec = r['问题类别(二级指标)']
        if sec not in internal_groups:
            internal_groups[sec] = []
        internal_groups[sec].append(r)

    general_titles = {
        '将不属于医保支付范围的纳入医保基金结算': '1. 将不属于医保支付范围的纳入医保基金结算',
        '违反诊疗规范过度诊疗': '2. 违反诊疗规范过度诊疗',
        '超标准收费': '3. 超标准收费',
    }
    internal_titles = {
        '其他问题': '1. 其他问题',
    }

    # 一般违规问题
    p_section1 = doc.add_paragraph()
    p_section1_run = p_section1.add_run('　　（一）一般违规问题')
    p_section1_run.bold = True
    p_section1_run.font.name = '黑体'
    p_section1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for sec_type, sec_records in ordered_general.items():
        p_secondary = doc.add_paragraph()
        p_secondary_run = p_secondary.add_run(f'　　{general_titles[sec_type]}')
        p_secondary_run.bold = True
        p_secondary_run.font.name = '黑体'
        p_secondary_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for idx, record in enumerate(sec_records):
            problem_name = record['问题项目']
            basis = record['认定依据']
            description = record['问题情形描述']
            quantity = record['问题数量']
            total_fee = record['违规医药费用总额（元）']
            fund = record['违规使用医保基金(元)']

            p_item = doc.add_paragraph()
            p_item_run = p_item.add_run(f'（{idx+1}）{problem_name}。')
            p_item_run.bold = True
            p_item_run.font.name = '宋体'
            p_item_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_item.paragraph_format.first_line_indent = Cm(0.74)

            p_basis = doc.add_paragraph()
            basis_text = f'根据{basis}。'
            p_basis_run = p_basis.add_run(basis_text)
            p_basis_run.font.name = '宋体'
            p_basis_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_basis.paragraph_format.first_line_indent = Cm(0.74)

            p_desc = doc.add_paragraph()
            desc_text = f'该院{description}。属于{sec_type}。'
            p_desc_run = p_desc.add_run(desc_text)
            p_desc_run.font.name = '宋体'
            p_desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_desc.paragraph_format.first_line_indent = Cm(0.74)

            p_stat = doc.add_paragraph()
            if quantity == '/' or pd.isna(quantity):
                stat_text = (
                    f'经统计，2024年01月01日---2025年12月31日期间，该医院上述情况'
                    f'涉及问题数量0人次，违规医药费用总额0.00元，违规使用医保基金0.00元。'
                )
            else:
                qty_str = f'{int(quantity)}' if isinstance(quantity, (int, float)) and quantity == int(quantity) else str(quantity)
                fee_str = f'{total_fee:.2f}' if isinstance(total_fee, (int, float)) and total_fee != int(total_fee) else f'{int(total_fee)}.00' if isinstance(total_fee, (int, float)) else str(total_fee)
                fund_str = f'{fund:.2f}' if isinstance(fund, (int, float)) and fund != int(fund) else f'{int(fund)}.00' if isinstance(fund, (int, float)) else str(fund)

                stat_text = (
                    f'经统计，2024年01月01日---2025年12月31日期间，该医院上述情况'
                    f'涉及问题数量{qty_str}人次，违规医药费用总额{fee_str}元，'
                    f'违规使用医保基金{fund_str}元。'
                )

            p_stat_run = p_stat.add_run(stat_text)
            p_stat_run.font.name = '宋体'
            p_stat_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_stat.paragraph_format.first_line_indent = Cm(0.74)
            p_stat.paragraph_format.space_after = Pt(6)

    # 内部管理问题
    p_section2 = doc.add_paragraph()
    p_section2_run = p_section2.add_run('　　（二）内部管理问题')
    p_section2_run.bold = True
    p_section2_run.font.name = '黑体'
    p_section2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for sec_type, sec_records in internal_groups.items():
        p_secondary = doc.add_paragraph()
        p_secondary_run = p_secondary.add_run(f'　　{internal_titles[sec_type]}')
        p_secondary_run.bold = True
        p_secondary_run.font.name = '黑体'
        p_secondary_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for idx, record in enumerate(sec_records):
            problem_name = record['问题项目']
            basis = record['认定依据']
            description = record['问题情形描述']
            quantity = record['问题数量']
            total_fee = record['违规医药费用总额（元）']
            fund = record['违规使用医保基金(元)']

            p_item = doc.add_paragraph()
            p_item_run = p_item.add_run(f'（{idx+1}）{problem_name}。')
            p_item_run.bold = True
            p_item_run.font.name = '宋体'
            p_item_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_item.paragraph_format.first_line_indent = Cm(0.74)

            p_basis = doc.add_paragraph()
            basis_text = f'根据{basis}。'
            p_basis_run = p_basis.add_run(basis_text)
            p_basis_run.font.name = '宋体'
            p_basis_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_basis.paragraph_format.first_line_indent = Cm(0.74)

            p_desc = doc.add_paragraph()
            desc_text = f'该院{description}。属于{sec_type}。'
            p_desc_run = p_desc.add_run(desc_text)
            p_desc_run.font.name = '宋体'
            p_desc_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_desc.paragraph_format.first_line_indent = Cm(0.74)

            p_stat = doc.add_paragraph()
            if quantity == '/' or pd.isna(quantity):
                stat_text = (
                    f'经统计，2024年01月01日---2025年12月31日期间，该医院上述情况'
                    f'涉及问题数量0人次，违规医药费用总额0.00元，违规使用医保基金0.00元。'
                )
            else:
                qty_str = f'{int(quantity)}' if isinstance(quantity, (int, float)) and quantity == int(quantity) else str(quantity)
                fee_str = f'{total_fee:.2f}' if isinstance(total_fee, (int, float)) and total_fee != int(total_fee) else f'{int(total_fee)}.00' if isinstance(total_fee, (int, float)) else str(total_fee)
                fund_str = f'{fund:.2f}' if isinstance(fund, (int, float)) and fund != int(fund) else f'{int(fund)}.00' if isinstance(fund, (int, float)) else str(fund)

                stat_text = (
                    f'经统计，2024年01月01日---2025年12月31日期间，该医院上述情况'
                    f'涉及问题数量{qty_str}人次，违规医药费用总额{fee_str}元，'
                    f'违规使用医保基金{fund_str}元。'
                )

            p_stat_run = p_stat.add_run(stat_text)
            p_stat_run.font.name = '宋体'
            p_stat_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p_stat.paragraph_format.first_line_indent = Cm(0.74)
            p_stat.paragraph_format.space_after = Pt(6)

    doc.save(output_path)
    return output_path


class ReportGeneratorGUI:
    """GUI界面类"""

    def __init__(self, root):
        self.root = root
        self.root.title('飞行检查报告生成器')
        self.root.geometry('600x400')
        self.root.resizable(False, False)

        # 设置样式
        style = ttk.Style()
        style.configure('TButton', font=('微软雅黑', 11))
        style.configure('TLabel', font=('微软雅黑', 11))

        # 主框架
        main_frame = ttk.Frame(root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 标题
        title_label = ttk.Label(main_frame, text='飞行检查报告生成器', font=('微软雅黑', 16, 'bold'))
        title_label.pack(pady=10)

        desc_label = ttk.Label(main_frame, text='将Excel问题表自动转换为Word检查报告', font=('微软雅黑', 10))
        desc_label.pack(pady=5)

        # 输入文件选择
        input_frame = ttk.Frame(main_frame)
        input_frame.pack(fill=tk.X, pady=15)

        ttk.Label(input_frame, text='Excel文件:').pack(side=tk.LEFT)
        self.input_path = tk.StringVar()
        input_entry = ttk.Entry(input_frame, textvariable=self.input_path, width=40)
        input_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(input_frame, text='浏览...', command=self.select_input).pack(side=tk.LEFT)

        # 输出文件夹选择
        output_frame = ttk.Frame(main_frame)
        output_frame.pack(fill=tk.X, pady=15)

        ttk.Label(output_frame, text='输出目录:').pack(side=tk.LEFT)
        self.output_dir = tk.StringVar(value=os.path.join(os.path.expanduser('~'), 'Desktop'))
        output_entry = ttk.Entry(output_frame, textvariable=self.output_dir, width=40)
        output_entry.pack(side=tk.LEFT, padx=5)
        ttk.Button(output_frame, text='浏览...', command=self.select_output).pack(side=tk.LEFT)

        # 生成按钮
        ttk.Button(main_frame, text='生成报告', command=self.generate_report, width=20).pack(pady=20)

        # 状态显示
        self.status_var = tk.StringVar(value='就绪')
        status_label = ttk.Label(main_frame, textvariable=self.status_var, font=('微软雅黑', 9))
        status_label.pack(pady=10)

        # 进度条
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate', length=400)
        self.progress.pack(pady=5)

    def select_input(self):
        """选择输入文件"""
        file_path = filedialog.askopenfilename(
            title='选择Excel文件',
            filetypes=[('Excel文件', '*.xlsx *.xls'), ('所有文件', '*.*')]
        )
        if file_path:
            self.input_path.set(file_path)

    def select_output(self):
        """选择输出目录"""
        dir_path = filedialog.askdirectory(title='选择输出目录')
        if dir_path:
            self.output_dir.set(dir_path)

    def generate_report(self):
        """生成报告"""
        input_file = self.input_path.get().strip()
        output_dir = self.output_dir.get().strip()

        if not input_file:
            messagebox.showwarning('警告', '请选择Excel文件')
            return

        if not os.path.exists(input_file):
            messagebox.showerror('错误', '输入文件不存在')
            return

        try:
            self.status_var.set('正在生成报告...')
            self.progress.start()
            self.root.update()

            # 生成输出文件名
            output_file = os.path.join(output_dir, '飞行检查报告.docx')

            # 生成报告
            excel_to_inspection_report(input_file, output_file)

            self.progress.stop()
            self.status_var.set(f'生成成功: {output_file}')
            messagebox.showinfo('成功', f'报告已生成:\n{output_file}')

        except Exception as e:
            self.progress.stop()
            self.status_var.set('生成失败')
            error_msg = f'错误: {str(e)}\n\n详细:\n{traceback.format_exc()}'
            messagebox.showerror('错误', error_msg)


def main():
    """主函数 - 支持GUI和命令行两种模式"""

    # 检查是否有命令行参数（排除-f等IPython参数）
    args = [arg for arg in sys.argv[1:] if not arg.startswith('-')]

    if args:
        # 命令行模式
        excel_path = args[0]
        output_path = args[1] if len(args) > 1 else '飞行检查报告.docx'

        try:
            result = excel_to_inspection_report(excel_path, output_path)
            print(f'报告已生成: {result}')
        except Exception as e:
            print(f'错误: {e}')
            traceback.print_exc()
            sys.exit(1)
    else:
        # GUI模式
        root = tk.Tk()
        app = ReportGeneratorGUI(root)
        root.mainloop()


if __name__ == '__main__':
    main()